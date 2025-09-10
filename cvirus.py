import sys
import os
from PyQt5.QtWidgets import (
    QApplication, QWidget, QLabel, QPushButton, QTextEdit, QVBoxLayout, 
    QHBoxLayout, QGridLayout, QFileDialog, QComboBox
)
from PyQt5.QtGui import QFont
from PyQt5.QtCore import Qt

import docx
import fitz  # PyMuPDF for PDF reading
import requests  # For Gemini API
import json

# --- Gemini API Helper ---
GEMINI_API_KEY = "gemini_api_key"  # <-- Replace with your Gemini API key
GEMINI_API_URL = "https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash:generateContent?key=" + GEMINI_API_KEY

def call_gemini(prompt):
    headers = {"Content-Type": "application/json"}
    data = {
        "contents": [{"parts": [{"text": prompt}]}]
    }
    try:
        response = requests.post(GEMINI_API_URL, headers=headers, data=json.dumps(data))
        response.raise_for_status()
        result = response.json()
        # Extract the generated text
        return result["candidates"][0]["content"]["parts"][0]["text"]
    except Exception as e:
        return f"[Gemini API error: {e}]"

class ResumeBuilder(QWidget):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("AscendCV")
        self.setMinimumSize(1000, 600)
        self.resume_text = ""  # Will store extracted resume text
        self.job_text = ""     # Will store extracted job description text
        self.last_saved_file = None
        self.generated_resume_text = ""
        self.init_ui()

    def init_ui(self):
        # --- Header ---
        header_label = QLabel("AscendCV")
        header_label.setFont(QFont("Arial", 20, QFont.Bold))
        header_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        header_label.setStyleSheet("color: #2c3e50; padding: 10px;")

        # --- Upload Resume ---
        upload_resume_btn = QPushButton("Upload Resume (PDF/DOCX)")
        upload_resume_btn.clicked.connect(self.upload_resume)
        self.resume_file_label = QLabel("No resume uploaded")
        self.resume_file_label.setStyleSheet("color: gray; font-style: italic;")

        # --- Upload Job Description ---
        upload_jd_btn = QPushButton("Upload Job Description (PDF)")
        upload_jd_btn.clicked.connect(self.upload_job_description)
        self.jd_file_label = QLabel("No job description uploaded")
        self.jd_file_label.setStyleSheet("color: gray; font-style: italic;")

        # --- Job Description Input (optional) ---
        self.job_input = QTextEdit()
        self.job_input.setPlaceholderText("Paste the job description here or upload a PDF...")

        # --- Labels ---
        job_label = QLabel("Job Description:")
        preview_label = QLabel("Output Preview")
        preview_label.setAlignment(Qt.AlignmentFlag.AlignCenter)

        # --- Theme Dropdown ---
        self.theme_dropdown = QComboBox()
        self.theme_dropdown.addItems(["Modern", "Classic", "Creative", "Minimal"])

        # --- Buttons ---
        generate_button = QPushButton("Generate Resume")
        reset_button = QPushButton("Reset")
        save_pdf_button = QPushButton("Save as PDF")
        save_txt_button = QPushButton("Save as TXT")

        # --- Output Preview ---
        self.output_preview = QTextEdit()
        self.output_preview.setReadOnly(True)
        self.output_preview.setMinimumHeight(400)

        # --- Button Actions ---
        generate_button.clicked.connect(self.generate_resume)
        reset_button.clicked.connect(self.reset_fields)
        save_pdf_button.clicked.connect(self.save_pdf)
        save_txt_button.clicked.connect(self.save_txt)

        # --- Layouts ---
        grid = QGridLayout()
        grid.setSpacing(10)

        # Left Column: Upload + Job Description
        grid.addWidget(upload_resume_btn, 0, 0)
        grid.addWidget(self.resume_file_label, 1, 0)
        grid.addWidget(upload_jd_btn, 2, 0)
        grid.addWidget(self.jd_file_label, 3, 0)
        grid.addWidget(job_label, 4, 0)
        grid.addWidget(self.job_input, 5, 0)

        # Center Column: Theme + Buttons
        center_layout = QVBoxLayout()
        center_layout.addWidget(QLabel("Resume Theme:"))
        center_layout.addWidget(self.theme_dropdown)
        center_layout.addSpacing(20)
        center_layout.addWidget(generate_button)
        center_layout.addWidget(reset_button)
        center_widget = QWidget()
        center_widget.setLayout(center_layout)
        grid.addWidget(center_widget, 1, 1, 5, 1)

        # Right Column: Output
        grid.addWidget(preview_label, 0, 2)
        grid.addWidget(self.output_preview, 1, 2, 6, 1)
        grid.addWidget(save_pdf_button, 7, 2)
        grid.addWidget(save_txt_button, 8, 2)

        # Master Layout
        main_layout = QVBoxLayout()
        main_layout.addWidget(header_label)
        main_layout.addLayout(grid)

        self.setLayout(main_layout)
        self.setStyleSheet("""
            QTextEdit {
                font-family: 'Segoe UI';
                font-size: 14px;
                border: 1px solid #ccc;
                border-radius: 6px;
                padding: 8px;
                background-color: #fdfdfd;
            }
            QLabel {
                font-weight: bold;
                color: #34495e;
            }
            QPushButton {
                background-color: #3498db;
                color: white;
                border: none;
                padding: 8px 12px;
                border-radius: 4px;
            }
            QPushButton:hover {
                background-color: #2980b9;
            }
            QComboBox {
                font-size: 14px;
            }
        """)

    def upload_resume(self):
        file_path, _ = QFileDialog.getOpenFileName(self, "Open Resume", "", "Documents (*.pdf *.docx)")
        if not file_path:
            return
        ext = os.path.splitext(file_path)[1].lower()
        text = ""
        try:
            if ext == ".docx":
                doc = docx.Document(file_path)
                text = "\n".join([para.text for para in doc.paragraphs])
            elif ext == ".pdf":
                doc = fitz.open(file_path)
                for page in doc:
                    get_text_fn = getattr(page, 'get_text', None)
                    if callable(get_text_fn):
                        text += str(get_text_fn())
                    else:
                        getText_fn = getattr(page, 'getText', None)
                        if callable(getText_fn):
                            text += str(getText_fn())
                        else:
                            text += "[Cannot extract text from page]"
            else:
                text = "[Unsupported file type]"
        except Exception as e:
            text = f"[Error reading file: {e}]"
        self.resume_text = text
        self.resume_file_label.setText(f"Uploaded: {os.path.basename(file_path)}")

    def upload_job_description(self):
        file_path, _ = QFileDialog.getOpenFileName(self, "Open Job Description", "", "PDF Files (*.pdf)")
        if not file_path:
            return
        text = ""
        try:
            doc = fitz.open(file_path)
            for page in doc:
                get_text_fn = getattr(page, 'get_text', None)
                if callable(get_text_fn):
                    text += str(get_text_fn())
                else:
                    getText_fn = getattr(page, 'getText', None)
                    if callable(getText_fn):
                        text += str(getText_fn())
                    else:
                        text += "[Cannot extract text from page]"
        except Exception as e:
            text = f"[Error reading file: {e}]"
        self.job_text = text
        self.jd_file_label.setText(f"Uploaded: {os.path.basename(file_path)}")
        self.job_input.setPlainText(text)

    def generate_resume(self):
        resume = self.resume_text.strip()
        job = self.job_input.toPlainText().strip() or self.job_text.strip()
        theme = self.theme_dropdown.currentText()
        if not resume or not job:
            self.output_preview.setPlainText("Please upload a resume and job description.")
            return
        # --- Gemini Prompt ---
        prompt = f"""
You are an advanced ATS resume optimization assistant.
Thoroughly analyze the following job description and the uploaded resume. Create a new, top-quality, ATS-friendly resume that:
- PRESERVE ALL SECTION TITLES, HEADERS, AND FORMATTING from the uploaded resume. Do NOT change the template, layout, or section names.
- Only update the content within each section to better align with the job description, using information from the uploaded resume.
- Intelligently expand and enhance the content within each section, but do NOT add new sections or change section titles.
- Do NOT add any data or skills that are not present or implied in the uploaded resume unless the job description explicitly mentions them.
- Use the '{theme}' theme for tone and style if possible.
- THE FINAL RESUME CONTENT MUST BE STRICTLY BETWEEN 550 AND 950 WORDS. If the content is too short, expand it with relevant details from the resume. If too long, summarize and condense as needed.
- Minimize free spaces and ensure the formatting is compact, professional, and highly relevant for the specific job role.
- The primary goal is to maximize ATS compatibility and ensure the resume is highly likely to be considered for the job role.

---
Job Description:
{job}
---
Resume:
{resume}
---
Output (new resume only, with the SAME section titles and structure as the uploaded resume):
"""
        output = call_gemini(prompt)
        self.output_preview.setPlainText(output.strip())
        self.generated_resume_text = output.strip()

    def reset_fields(self):
        self.resume_text = ""
        self.job_text = ""
        self.job_input.clear()
        self.output_preview.clear()
        self.resume_file_label.setText("No resume uploaded")
        self.jd_file_label.setText("No job description uploaded")
        self.theme_dropdown.setCurrentIndex(0)
        self.last_saved_file = None
        self.generated_resume_text = ""

    def save_pdf(self):
        file_name, _ = QFileDialog.getSaveFileName(self, "Save Resume as PDF", "", "PDF Files (*.pdf)")
        if file_name:
            # Use QTextDocument to export to PDF
            from PyQt5.QtGui import QTextDocument
            from PyQt5.QtPrintSupport import QPrinter
            doc = QTextDocument()
            doc.setPlainText(self.output_preview.toPlainText())
            printer = QPrinter()
            printer.setOutputFormat(QPrinter.PdfFormat)
            printer.setOutputFileName(file_name)
            doc.print_(printer)
            self.last_saved_file = file_name
            self.show_google_docs_link(file_name)

    def save_txt(self):
        file_name, _ = QFileDialog.getSaveFileName(self, "Save Resume as TXT", "", "Text Files (*.txt)")
        if file_name:
            with open(file_name, 'w', encoding='utf-8') as f:
                f.write(self.output_preview.toPlainText())
            self.last_saved_file = file_name
            self.show_google_docs_link(file_name)

    def show_google_docs_link(self, file_path):
        # Google Docs upload URL (user must upload manually, but we can provide the link)
        import webbrowser
        from PyQt5.QtWidgets import QMessageBox
        msg = QMessageBox()
        msg.setWindowTitle("Open in Google Docs")
        msg.setText("To open your resume in Google Docs, click the button below to go to the Google Docs upload page. Then upload your saved file.")
        msg.setStandardButtons(QMessageBox.Open | QMessageBox.Close)
        ret = msg.exec_()
        if ret == QMessageBox.Open:
            webbrowser.open("https://docs.google.com/document/u/0/?usp=docs_home&ths=true")

if __name__ == "__main__":
    app = QApplication(sys.argv)
    win = ResumeBuilder()
    win.show()
    sys.exit(app.exec_())
